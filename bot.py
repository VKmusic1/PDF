import os
import io
import logging
import time
import asyncio
import threading

import fitz                  # PyMuPDF
import pdfplumber
import pandas as pd
from flask import Flask, request
from telegram import (
    Update,
    InputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup
)
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters
)
from docx import Document
from pdf2docx import Converter
import requests

# ---------------------- 1. Конфигурация из окружения ----------------------
TOKEN = os.getenv("TOKEN")
HOST  = os.getenv("RENDER_EXTERNAL_HOSTNAME")
PORT  = int(os.getenv("PORT", "10000"))
if not TOKEN or not HOST:
    raise RuntimeError("Environment variables TOKEN and RENDER_EXTERNAL_HOSTNAME are required")
WEBHOOK_URL = f"https://{HOST}/{TOKEN}"

# Разбиение на куски по N страниц
CHUNK_SIZE = 30

# ---------------------- 2. Логирование ----------------------
logging.basicConfig(format="%(asctime)s %(levelname)s: %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------- 3. Flask ----------------------
app = Flask(__name__)

# ---------------------- 4. asyncio-loop ----------------------
telegram_loop = asyncio.new_event_loop()
def start_loop(loop):
    asyncio.set_event_loop(loop)
    loop.run_forever()
threading.Thread(target=start_loop, args=(telegram_loop,), daemon=True).start()

# ---------------------- 5. Инициализация Telegram ----------------------
telegram_app = (
    Application.builder()
    .token(TOKEN)
    .connection_pool_size(100)
    .build()
)
# HTTP таймауты
telegram_app.request_kwargs = {"read_timeout":60,"connect_timeout":20}
# Инициализируем PTB
future = asyncio.run_coroutine_threadsafe(telegram_app.initialize(), telegram_loop)
future.result(timeout=15)
logger.info("✔ Telegram application initialized")

# ---------------------- 6. PDF-утилиты ----------------------
def extract_pdf_pages(path: str):
    """Возвращает список по страницам, где каждая страница это список элементов (тип, контент)."""
    doc = fitz.open(path)
    pages = []
    for page in doc:
        elems = []
        text = page.get_text().strip()
        if text:
            elems.append(("text", text))
        for img in page.get_images(full=True):
            data = doc.extract_image(img[0])["image"]
            elems.append(("img", data))
        pages.append(elems)
    doc.close()
    return pages

def convert_layout(pdf_path: str, out_path: str):
    """Конвертация макета через pdf2docx."""
    conv = Converter(pdf_path)
    conv.convert(out_path, start=0, end=None)
    conv.close()

def convert_to_word(elements, out_path: str):
    """Просто текст+картинки в Word."""
    docx = Document()
    for typ, cnt in elements:
        if typ == "text":
            docx.add_paragraph(cnt)
        else:
            bio = io.BytesIO(cnt)
            bio.name = "image.png"
            docx.add_picture(bio)
    docx.save(out_path)

# ---------------------- 7. Хендлеры ----------------------
async def start(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.message.reply_text("👋 Привет! Пришли мне PDF-файл, и я предложу варианты извлечения.")

async def handle_pdf(u: Update, c: ContextTypes.DEFAULT_TYPE):
    doc = u.message.document
    if not doc or doc.mime_type != "application/pdf":
        return await u.message.reply_text("❌ Пожалуйста, отправь PDF.")
    # скачиваем
    tgfile = await doc.get_file()
    pdf_path = f"/tmp/{doc.file_unique_id}.pdf"
    await tgfile.download_to_drive(pdf_path)
    # разбиваем на страницы
    pages = extract_pdf_pages(pdf_path)
    c.user_data["pdf_path"] = pdf_path
    c.user_data["pages"] = pages
    # клавиатура
    kb = [
        [InlineKeyboardButton("Word: макет 📄", callback_data="cb_layout")],
        [InlineKeyboardButton("Word: текст+картинки 📝", callback_data="cb_all")],
        [InlineKeyboardButton("TXT: текст 📄", callback_data="cb_txt")],
        [InlineKeyboardButton("Excel: таблицы 📊", callback_data="cb_tables")],
        [InlineKeyboardButton("Чат: текст 📝", callback_data="cb_text")],
        [InlineKeyboardButton("Чат: текст+картинки 🖼️📝", callback_data="cb_chat")],
        [InlineKeyboardButton("Новый PDF 🔄", callback_data="cb_new")]
    ]
    await u.message.reply_text("Выбери вариант обработки:", reply_markup=InlineKeyboardMarkup(kb))

async def cb_layout(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    path = c.user_data.get("pdf_path")
    if not path:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет PDF — пришли заново.")
    out = f"/tmp/{u.effective_user.id}_layout.docx"
    msg = await c.bot.send_message(u.effective_chat.id, "⏳ Конвертация макета...")
    # конверсия
    convert_layout(path, out)
    await c.bot.edit_message_text("✅ Готово!", u.effective_chat.id, msg.message_id)
    with open(out, "rb") as f:
        await c.bot.send_document(u.effective_chat.id, InputFile(f, filename="layout.docx"))

async def cb_all(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    pages = c.user_data.get("pages", [])
    if not pages:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет контента.")
    # собираем все элементы
    elems = [elem for pg in pages for elem in pg]
    out = f"/tmp/{u.effective_user.id}_all.docx"
    convert_to_word(elems, out)
    with open(out, "rb") as f:
        await c.bot.send_document(u.effective_chat.id, InputFile(f, filename="all.docx"))

async def cb_txt(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    pages = c.user_data.get("pages", [])
    text_blocks = [txt for pg in pages for typ, txt in pg if typ == "text"]
    if not text_blocks:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет текста.")
    all_text = "\n\n".join(text_blocks)
    out = f"/tmp/{u.effective_user.id}.txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write(all_text)
    with open(out, "rb") as f:
        await c.bot.send_document(u.effective_chat.id, InputFile(f, filename="text.txt"))

async def cb_tables(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    path = c.user_data.get("pdf_path")
    if not path:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет PDF.")
    all_tables = []
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        for i, pg in enumerate(pdf.pages, 1):
            for ti, tbl in enumerate(pg.extract_tables(), 1):
                df = pd.DataFrame(tbl[1:], columns=tbl[0])
                all_tables.append((f"Page{i}_Tbl{ti}", df))
    if not all_tables:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет таблиц.")
    xlsx = f"/tmp/{u.effective_user.id}_tabs.xlsx"
    with pd.ExcelWriter(xlsx, engine="openpyxl") as w:
        for name, df in all_tables:
            df.to_excel(w, sheet_name=name[:31], index=False)
    with open(xlsx, "rb") as f:
        await c.bot.send_document(u.effective_chat.id, InputFile(f, filename="tables.xlsx"))

async def cb_text(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    pages = c.user_data.get("pages", [])
    if not pages:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет контента.")
    # разбиваем на чанки по CHUNK_SIZE страниц
    total = len(pages)
    for start in range(0, total, CHUNK_SIZE):
        end = min(start + CHUNK_SIZE, total)
        header = f"📝 Отправляю страницы {start+1}–{end}/{total}"
        await c.bot.send_message(u.effective_chat.id, header)
        for pg in pages[start:end]:
            for typ, cnt in pg:
                if typ == "text":
                    for i in range(0, len(cnt), 4096):
                        await c.bot.send_message(u.effective_chat.id, cnt[i:i+4096])
        await asyncio.sleep(0.5)

async def cb_chat(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    pages = c.user_data.get("pages", [])
    if not pages:
        return await c.bot.send_message(u.effective_chat.id, "❌ Нет контента.")
    total = len(pages)
    for start in range(0, total, CHUNK_SIZE):
        end = min(start + CHUNK_SIZE, total)
        header = f"🖼️📝 Страницы {start+1}–{end}/{total}"
        await c.bot.send_message(u.effective_chat.id, header)
        sent = set()
        for pg in pages[start:end]:
            for typ, cnt in pg:
                if typ == "text":
                    for i in range(0, len(cnt), 4096):
                        await c.bot.send_message(u.effective_chat.id, cnt[i:i+4096])
                else:
                    h = hash(cnt)
                    if h in sent:
                        continue
                    sent.add(h)
                    bio = io.BytesIO(cnt)
                    bio.name = "image.png"
                    await c.bot.send_photo(u.effective_chat.id, photo=bio)
        await asyncio.sleep(0.5)

async def cb_new(u: Update, c: ContextTypes.DEFAULT_TYPE):
    await u.callback_query.answer()
    c.user_data.clear()
    await c.bot.send_message(u.effective_chat.id, "🔄 Готов к новому PDF!")

# регистрация
telegram_app.add_handler(CommandHandler("start", start))
telegram_app.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))
telegram_app.add_handler(CallbackQueryHandler(cb_layout, pattern="cb_layout"))
telegram_app.add_handler(CallbackQueryHandler(cb_all,    pattern="cb_all"))
telegram_app.add_handler(CallbackQueryHandler(cb_txt,    pattern="cb_txt"))
telegram_app.add_handler(CallbackQueryHandler(cb_tables, pattern="cb_tables"))
telegram_app.add_handler(CallbackQueryHandler(cb_text,   pattern="cb_text"))
telegram_app.add_handler(CallbackQueryHandler(cb_chat,   pattern="cb_chat"))
telegram_app.add_handler(CallbackQueryHandler(cb_new,    pattern="cb_new"))

# ---------------------- 8. Webhook Routes ----------------------
@app.route(f"/{TOKEN}", methods=["POST"])
def telegram_webhook():
    data = request.get_json(force=True)
    upd = Update.de_json(data, telegram_app.bot)
    asyncio.run_coroutine_threadsafe(telegram_app.process_update(upd), telegram_loop)
    return "OK"

@app.route("/ping")
def ping():
    return "pong"

# ---------------------- 9. Запуск ----------------------
if __name__ == "__main__":
    logger.info("Setting webhook to %s", WEBHOOK_URL)
    r = requests.post(f"https://api.telegram.org/bot{TOKEN}/setWebhook", data={"url": WEBHOOK_URL})
    if not r.ok:
        logger.error("Webhook error: %s", r.text)
    logger.info("Running Flask on port %s", PORT)
    app.run(host="0.0.0.0", port=PORT)
